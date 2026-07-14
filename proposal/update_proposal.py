"""
Updates proposal/proposal.docx to reflect work done since the original draft:
the calibrated analytical model (v2) and the ChampSim empirical validation
(Phase 0-2, through v4). Adds a new "Empirical Validation in ChampSim"
section with embedded figures (see generate_figures.py), rewrites the
analytical-model subsection to use the calibrated numbers instead of the
original uncalibrated spike's, and renumbers downstream sections.

Run generate_figures.py first. Then:  python3 proposal/update_proposal.py
"""

import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

P_TAG = qn("w:p")


def next_paragraph(p):
    """Like p._element.getnext() but skips non-<w:p> siblings and returns
    a Paragraph, or None if there isn't another paragraph sibling."""
    el = p._element.getnext()
    while el is not None and el.tag != P_TAG:
        el = el.getnext()
    return Paragraph(el, p._parent) if el is not None else None

HERE = os.path.dirname(__file__)
DOCX_PATH = os.path.join(HERE, "proposal.docx")
FIG_DIR = os.path.join(HERE, "figures")
IMG_WIDTH = Inches(5.8)
MUTED = RGBColor(0x52, 0x51, 0x4E)

doc = docx.Document(DOCX_PATH)


# ------------------------------------------------------------- utilities --
def find_para(text_exact=None, text_startswith=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith!r}")


def delete_paragraph(p):
    p._element.getparent().remove(p._element)


def rename_heading_prefix(old_prefix, new_prefix):
    p = find_para(text_startswith=old_prefix)
    rest = p.text[len(old_prefix):]
    new_text = new_prefix + rest
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text
    return p


def insert_heading_before(anchor, text, level):
    return anchor.insert_paragraph_before(text, style=f"Heading {level}")


def insert_para_before(anchor, text, style="Normal"):
    return anchor.insert_paragraph_before(text, style=style)


def insert_image_before(anchor, filename, caption):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(FIG_DIR, filename), width=IMG_WIDTH)
    cap = anchor.insert_paragraph_before(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.color.rgb = MUTED


# =========================================================== 1. status line
status_p = find_para(text_startswith="Run ID:")
status_p.runs[0].text = "Run ID: 20260409_140834_run_010  |  Category: Memory Systems\n"
status_p.runs[1].text = ("Status: Literature-vetted, analytically validated (calibrated model), and "
                          "empirically validated in ChampSim (4 of 5 tracked load sites, 429.mcf)")


# ============================================ 2. rewrite the 2.4 subsection
# Delete the old "2.4 Validation Spike ..." subsection (heading through its
# closing paragraph), which cited the uncalibrated v1 spike's numbers
# (0.4%-40.1% wasted bandwidth) -- since audited and superseded by the
# calibrated v2 model in analytical_model/.
old_heading = find_para(text_startswith="2.4 Validation Spike")
old_paras = [old_heading]
# collect every paragraph from the heading up to (not including) "3. Approach"
p = old_heading
while True:
    np = next_paragraph(p)
    if np is None or np.text.strip().startswith("3. Approach"):
        break
    old_paras.append(np)
    p = np

approach_heading = find_para(text_startswith="3. Approach")  # keep a live anchor before deleting

for op in old_paras:
    delete_paragraph(op)

# Insert the new, calibrated 2.4 in its place (right before "3. Approach").
h = insert_heading_before(approach_heading, "2.4 Analytical Model: From Uncalibrated Spike to a Calibrated Model", 2)

insert_para_before(approach_heading,
    "Before committing to simulator time, a Python analytical model was built to put numbers on three "
    "questions: is the wrong-path prefetch rate itself large, does it translate into real DRAM bandwidth "
    "waste, and does today's existing reactive throttling already recover most of it. A first spike answered "
    "all three, but an independent audit of that spike found two real gaps: its bandwidth figure was derived "
    "from unanchored assumptions rather than a measured reference point, and its reactive-throttling recovery "
    "curve was asserted (a hand-tuned decay) rather than derived from anything.")

insert_image_before(approach_heading, "fig1_calibration_fix.png",
    "Figure 1. The uncalibrated spike's worst-case estimate was ~4x higher than what a real system of this "
    "class actually measures (Magellan, ISCA'25, Fig. 19). The calibrated model anchors to that measurement "
    "and can never repeat the error.")

insert_para_before(approach_heading,
    "The calibrated model (v2) fixes both gaps: total prefetch-bandwidth overhead is anchored to Magellan's "
    "own measured ~10% total overhead (swept +/-50% for other workloads/configs) instead of derived from "
    "scratch, and reactive-throttling recovery is derived from a Zipf/power-law model of how wrong-path "
    "waste is distributed across static prefetch sites, with a concentration parameter (alpha) that is "
    "independent of misprediction rate -- fixing the earlier model's conflation of the two. The three stages "
    "are then composed into one end-to-end number: wrong-path rate x wasted bandwidth x (1 - recovery), the "
    "single figure that actually matters -- bandwidth a proactive scheme would still need to recover.")

insert_image_before(approach_heading, "fig2_composed_sweep.png",
    "Figure 2. Composed residual bandwidth across representative parameter points. Typical cases land at "
    "0.1%-2.0% of one DRAM channel; the worst swept combination (elevated MPKI, diffuse waste) reaches 7.7%. "
    "This is an order of magnitude below the uncalibrated spike's 40% headline, and consistent with Magellan's "
    "own reported overhead.")

insert_para_before(approach_heading,
    "The strongest quantitative argument for the idea, however, is not the raw bandwidth figure but its "
    "translation into queueing-delay terms: the same residual bandwidth buys a disproportionately larger "
    "tail-latency reduction as baseline channel contention rises, since queueing delay scales "
    "super-linearly with utilization.")

insert_image_before(approach_heading, "fig3_latency_uplift.png",
    "Figure 3. The same residual wasted bandwidth produces a larger queueing-delay-proxy reduction at higher "
    "baseline channel utilization (rho_base) -- the mechanism's payoff concentrates exactly where the problem "
    "statement's own 99th-percentile-latency figure of merit would be measured, not spread evenly across all "
    "operating points.")

insert_para_before(approach_heading,
    "The sensitivity pass over the composed model ranks gating_branches (how many control-dependent branches "
    "actually gate a given prefetch's usefulness) as the highest-leverage unmeasured parameter, ahead of "
    "misprediction rate itself -- the concrete motivation for the ChampSim instrumentation pass described in "
    "Section 3.")


# ============================ 3. renumber "3. Approach" -> "4. Approach" etc
# (must rename in an order that doesn't create duplicate prefixes mid-way,
# so do these before inserting the new Section 3 content, using the
# now-unambiguous full old headings.)
rename_heading_prefix("3.3 Design Space for the Gate", "4.3 Design Space for the Gate")
rename_heading_prefix("3.2 Mechanism Sketch", "4.2 Mechanism Sketch")
rename_heading_prefix("3.1 Core Idea", "4.1 Core Idea")
rename_heading_prefix("4.5 Recommended Path", "5.5 Recommended Path")
rename_heading_prefix("4.4 Solution D", "5.4 Solution D")
rename_heading_prefix("4.3 Solution C", "5.3 Solution C")
rename_heading_prefix("4.2 Solution B", "5.2 Solution B")
rename_heading_prefix("4.1 Solution A", "5.1 Solution A")
rename_heading_prefix("4. Possible Solutions", "5. Possible Solutions")
rename_heading_prefix("5. References", "6. References")
approach_heading = rename_heading_prefix("3. Approach", "4. Approach")  # refresh anchor after edits


# ========================= 4. insert new Section 3: Empirical Validation ==
insert_heading_before(approach_heading, "3. Empirical Validation in ChampSim", 1)

insert_heading_before(approach_heading, "3.1 Methodology", 2)
insert_para_before(approach_heading,
    "The analytical model's two remaining unmeasured assumptions -- gating_branches and the waste-"
    "concentration parameter alpha -- were instrumented directly in ChampSim rather than continuing to sweep "
    "them. A simplified, Magellan-inspired prefetcher module (\"loop_guided\") was built for ChampSim: since "
    "ChampSim's trace format records only addresses and never the register/memory values a genuine indirect-"
    "address computation (A[B[i]]) would need, the module instead detects periodic address-delta patterns per "
    "load PC -- a proxy for \"this load revisits the same relative pattern every loop iteration\", the same "
    "structural property Magellan's own dependence-graph extraction exploits.")
insert_para_before(approach_heading,
    "A second finding shaped the instrumentation design: ChampSim's core does not model wrong-path "
    "instruction fetch at all -- every branch's correctness is checked against the trace's oracle-known "
    "outcome the instant it is fetched, so no instructions are ever fetched past a to-be-mispredicted branch. "
    "Reproducing genuine wrong-path prefetch dispatch was therefore not attempted; instead, wrong-path-"
    "equivalence is measured by asking whether the branch a prefetch's target implicitly depended on was "
    "mispredicted by the time the corresponding real access happened -- exactly the mechanism Magellan's own "
    "paper describes (Sec. 4.4) for why mispredicted loop induction variables corrupt prefetch targets. This "
    "only requires real per-branch outcomes and real dynamic timing, both of which ChampSim already models "
    "correctly.")

insert_heading_before(approach_heading, "3.2 Iterative Measurement Refinement", 2)
insert_para_before(approach_heading,
    "A code review across 8 finder passes and 3 independent verifier passes caught a series of real "
    "correctness gaps in the first instrumentation attempt, each fixed and re-measured:")
insert_para_before(approach_heading,
    "v1 counted every conditional branch retiring between a prefetch's issue and its use, not just ones "
    "control-dependent on that specific prefetch -- 87.8% of matched prefetches hit a 32-branch overflow cap. "
    "v2 scoped counting to each PC's most-likely gating branch, reducing the overflow bucket to 39% but "
    "revealing low identification confidence (31-45%). v3 fixed a warmup-boundary bug (the dominant remaining "
    "cause of inflated waste), bounded candidate identification to a coarse same-function IP-distance proxy, "
    "locked each PC's identified gating branch for consistency, and made every remaining bias (queue "
    "eviction, staleness, sentinel collisions) either fixed or a counted, visible CSV column instead of a "
    "silent one -- reducing the overflow bucket to 2.0%. v4 root-caused and fixed a Phase 1 prefetcher design "
    "issue the v3 fixes had only diagnosed: a single lookahead slot shared across all tracked load sites was "
    "repeatedly stomped and reset, issuing prefetches far faster than real occurrences could close them out.",
    style="Normal")

insert_image_before(approach_heading, "fig5_debugging_journey.png",
    "Figure 4. Overflow-bucket concentration (share of matched prefetches whose control-dependent-branch "
    "count exceeded the histogram's 32-branch cap) across three rounds of instrumentation fixes -- evidence "
    "of methodological convergence, not just a single uncontested measurement.")

insert_heading_before(approach_heading, "3.3 Results", 2)
insert_para_before(approach_heading,
    "After the v4 fix (per-PC lookahead state, plus a disclosed IPC-vs-measurement-reliability tradeoff "
    "confirmed with the project owner), 4 of the 5 tracked load sites in 429.mcf produce internally-"
    "consistent, trustworthy waste measurements -- validated as purely additive to simulated behavior "
    "throughout (the instrumentation never changes IPC or prefetch counts; only the deliberate v4 prefetcher-"
    "policy change did, disclosed separately).")

insert_image_before(approach_heading, "fig4_champsim_empirical.png",
    "Figure 5. Measured wrong-path-equivalent waste fraction per tracked load site. Two sites carry large "
    "sample sizes (~9,500 matched prefetches each) at an 18.3% waste rate -- inside the calibrated analytical "
    "model's predicted range (Figure 2). The fifth site remains excluded: its low absolute occurrence count "
    "suggests a genuinely low-frequency or irregularly-spaced access pattern the \"match against the next "
    "occurrence\" model doesn't suit, rather than a fixable measurement bias.")

insert_heading_before(approach_heading, "3.4 What This Does and Does Not Establish", 2)
insert_para_before(approach_heading,
    "This is real, corroborating empirical evidence for the analytical model's central claim -- two "
    "independent, well-diagnosed measurements landing inside the predicted range, not merely \"plausible in "
    "principle.\" It is not yet sufficient to bulk-replace the model's swept assumptions: the evidence comes "
    "from a single workload (mcf) over a short (5M-instruction) simulation window, and gating-branch "
    "identification remains a coarse IP-distance heuristic rather than genuine control-dependence analysis. "
    "The immediate next step is feeding these four real distributions into the analytical model in place of "
    "the swept GATING_BRANCHES_OPTIONS/ZIPF_ALPHA_OPTIONS ranges and checking whether the composed "
    "residual-bandwidth conclusion survives, followed by validation on 1-2 additional workloads, before any "
    "further investment in the mechanism proposed in Section 4.")


# ================================================ 5. add ChampSim reference
refs_heading = find_para(text_startswith="6. References")
p = refs_heading
last_ref = refs_heading
while True:
    np = next_paragraph(p)
    if np is None or not np.text.strip():
        break
    last_ref = np
    p = np

new_ref = last_ref.insert_paragraph_before(
    "[7] Gober, N., Chacon, G., Wang, L., Gratz, P. V., Jimenez, D. A., Teran, E., Pugsley, S., & Kim, J. "
    "(2022). The Championship Simulator: Architectural Simulation for Education and Competition. "
    "arXiv:2210.14324. (ChampSim, used for the empirical validation in Section 3.)",
    style="Normal")
# copy formatting from an existing reference paragraph if possible
if last_ref.runs and new_ref.runs:
    new_ref.runs[0].font.size = last_ref.runs[0].font.size

doc.save(DOCX_PATH)
print(f"Saved {DOCX_PATH}")
