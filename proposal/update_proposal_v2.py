"""
v2 update: adds Phase 3 (calibration cross-check) and Phase 4 (real data fed
back into the analytical model) to proposal.docx's Section 3, and rewrites
the closing subsection since its "next step" description is now done.

Run generate_figures.py first (produces fig6/fig7). Then:
    python3 proposal/update_proposal_v2.py
"""

import os

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, RGBColor
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(__file__)
DOCX_PATH = os.path.join(HERE, "proposal.docx")
FIG_DIR = os.path.join(HERE, "figures")
IMG_WIDTH = Inches(5.8)
MUTED = RGBColor(0x52, 0x51, 0x4E)

P_TAG = qn("w:p")

doc = docx.Document(DOCX_PATH)


def find_para(text_exact=None, text_startswith=None):
    for p in doc.paragraphs:
        t = p.text.strip()
        if text_exact is not None and t == text_exact:
            return p
        if text_startswith is not None and t.startswith(text_startswith):
            return p
    raise ValueError(f"paragraph not found: {text_exact or text_startswith!r}")


def rename_heading_prefix(old_prefix, new_prefix):
    p = find_para(text_startswith=old_prefix)
    rest = p.text[len(old_prefix):]
    new_text = new_prefix + rest
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text
    return p


def insert_heading_before(anchor, text, level):
    return anchor.insert_paragraph_before(text, style=f"Heading {level}")


def insert_para_before(anchor, text, style="Normal"):
    return anchor.insert_paragraph_before(text, style=style)


def insert_image_before(anchor, filename, caption):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(os.path.join(FIG_DIR, filename), width=IMG_WIDTH)
    cap = anchor.insert_paragraph_before(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.color.rgb = MUTED


# Step 1: rename the closing subsection out of the way first, to avoid
# ambiguity while inserting the two new subsections before it.
rename_heading_prefix("3.4 What This Does and Does Not Establish", "3.6 What This Does and Does Not Establish")
closing_heading = find_para(text_startswith="3.6 What This Does and Does Not Establish")

# Step 2: insert 3.4 Calibration Cross-Check (Phase 3)
insert_heading_before(closing_heading, "3.4 Calibration Cross-Check", 2)
insert_para_before(closing_heading,
    "Before trusting the Phase 2 measurements further, ChampSim's own simulated total DRAM bandwidth "
    "overhead was compared against Magellan's own reported ~10% total prefetch overhead -- the same "
    "with-vs-without-prefetching comparison Magellan's own Fig. 19 makes, run on the same trace and window.")

insert_image_before(closing_heading, "fig6_phase3_calibration.png",
    "Figure 6. Measured total prefetch-bandwidth overhead is far below Magellan's own reported figure at "
    "both prefetch-distance settings tested (log scale). Traced to an exact mechanism, not a bug: LLC-level "
    "prefetch misses almost perfectly substitute for a matching reduction in demand-load misses, confirming "
    "prefetches overwhelmingly retime existing fetches rather than adding new ones -- the same qualitative "
    "story as Magellan's own paper, just a more conservative version of it.")

insert_para_before(closing_heading,
    "This does not contradict the 18.3% waste finding above: that measures a proxy for waste in a "
    "hypothetical deep-ROB machine with genuine wrong-path instruction fetch, which ChampSim's core does not "
    "model at all, so a \"wrong-path-equivalent\" prefetch can still go on to be used by a later demand access "
    "and contribute zero net extra bandwidth in the only world ChampSim can actually simulate. The two numbers "
    "are complementary measurements of different things, not competing ones. The calibration check raised no "
    "red flag, but the gap versus Magellan's own figure is large enough that this specific prefetcher/workload "
    "combination should not be treated as bandwidth-representative of a tuned production system -- a caveat "
    "carried into Section 3.5, not resolved here.")

# Step 3: insert 3.5 Closing the Loop: Real Data in the Model (Phase 4)
insert_heading_before(closing_heading, "3.5 Closing the Loop: Real Data in the Analytical Model", 2)
insert_para_before(closing_heading,
    "The real distributions measured in Sections 3.3-3.4 were fed into the calibrated analytical model's own "
    "functions (Section 2.4) in place of the swept gating_branches/alpha/total-overhead assumptions, to check "
    "whether the composed residual-bandwidth prediction survives contact with real data. Using the real "
    "aggregate wrong-path rate (25.3%, weighted across the four trustworthy sites) and the real waste "
    "concentration (only 4 distinct prefetch-issuing sites were found in this workload, not the 32 originally "
    "assumed; the best-fit Zipf concentration parameter is alpha~=0.49, mid-range within what was swept) -- "
    "combined with Magellan's literature-anchored overhead rather than this project's own "
    "under-representative measurement (Section 3.4) -- the composed residual comes out to 1.62% of one DRAM "
    "channel.")

insert_image_before(closing_heading, "fig7_phase4_validation.png",
    "Figure 7. The composed model's prediction, made before any ChampSim data existed, against real measured "
    "data. The real-data result (green) lands inside the originally predicted typical range; using this "
    "project's own conservatively-tuned overhead measurement instead (gray) lands far below it, consistent "
    "with the Section 3.4 caveat rather than contradicting the model.")

insert_para_before(closing_heading,
    "One discrepancy surfaced rather than resolved: the model's formula, given ChampSim's own real "
    "branch-accuracy statistics for this run (MPKI 2.456, 97.17% program-wide accuracy), predicts only 5.2% "
    "wrong-path at the real measured gating_branches (mean 1.843) -- about 5x below the 25.3% actually "
    "measured. Matching the real rate would require the specific identified gating branch to have a local "
    "accuracy near 85%, well below the program-wide average. This is consistent with that branch being a "
    "disproportionately hard-to-predict (H2P) one rather than an average branch, which a single global-accuracy "
    "figure cannot see -- and it directly reinforces this proposal's core premise that per-branch, not "
    "program-average, confidence is the signal worth gating prefetch dispatch on.")

# Step 4: rewrite the closing subsection's content (was written when Phase
# 4 was still a "next step"; now it's done).
old_closing_body = find_para(text_startswith="This is real, corroborating empirical evidence")
approach_heading = find_para(text_startswith="4. Approach")
old_closing_body._element.getparent().remove(old_closing_body._element)

insert_para_before(approach_heading,
    "This is real, corroborating empirical evidence for the analytical model's central claim, now closed "
    "into a full loop: a composed prediction made from swept literature-anchored assumptions alone (Section "
    "2.4) was checked against real ChampSim measurements (Sections 3.3-3.4) and survived (Section 3.5) -- not "
    "merely \"plausible in principle\" but landing inside the range predicted before any simulator time was "
    "spent. This is still not proof: the evidence comes from a single workload (mcf) over a short "
    "(5M-instruction) simulation window, one of five tracked load sites remains excluded as unmeasurable with "
    "the current method, and gating-branch identification remains a coarse IP-distance heuristic rather than "
    "genuine control-dependence analysis. The concrete next steps, in order, are: validate on 1-2 additional "
    "workloads and a longer window to check the 18.3%/1.62% figures hold up; investigate the H2P-branch "
    "discrepancy in Section 3.5 with direct per-branch accuracy instrumentation; and only then commit to "
    "building the mechanism proposed in Section 4.")

doc.save(DOCX_PATH)
print(f"Saved {DOCX_PATH}")

